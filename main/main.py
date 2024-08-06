import tkinter as tk
from tkinter import ttk
from docx import Document
import datetime as dt

root = tk.Tk()
root.title("WriteAlign")
root.configure(background="white")
root.minsize(700, 500)
essay_name_var = tk.StringVar()
thesis_var = tk.StringVar()
topic_sentence_1_var = tk.StringVar()
topic_sentence_2_var = tk.StringVar()
topic_sentence_3_var = tk.StringVar()
var_1 = tk.StringVar()
var_2 = tk.StringVar()
var_3 = tk.StringVar()
var_4 = tk.StringVar()
var_5 = tk.StringVar()
var_6 = tk.StringVar()
summary_var = tk.StringVar()
var_7 = tk.StringVar()
var_8 = tk.StringVar()
var_9 = tk.StringVar()
var_10 = tk.StringVar()
var_11 = tk.StringVar()
var_12 = tk.StringVar()
var_13 = tk.StringVar()
var_14 = tk.StringVar()
file_name_var = tk.StringVar()


main_page = tk.Frame(root, background="white")
main_page.pack(side="top", fill='both', padx=150, pady=150)
main_page.columnconfigure(0, weight=1)
write_align_title = tk.Label(main_page, text="WriteAlign", font=("Times New Roman", 50, "bold"), background="white")
write_align_title.grid(row=0, column=0, pady=(0, 30))
template_var = tk.StringVar()
template_list = ["Simple Essay", "Simple Biography", "Simple Resume", "Simple Cover Letter"]
select_template = ttk.Combobox(main_page, textvariable=template_var, values=template_list, state="readonly", font=("Times New Roman", 16), justify="center")
select_template.grid(row=1, column=0)
template_var.set(template_list[0])
finish_frame = tk.Frame()
general_frame = tk.Frame(root)


# Template Frames
simple_essay_frame = tk.Frame(root)

label_font = ("Times New Roman", 16)
entry_font = ("Times New Roman", 12)


def reset_variables(*args):
    essay_name_var.set("")
    thesis_var.set("")
    topic_sentence_1_var.set("")
    topic_sentence_2_var.set("")
    topic_sentence_3_var.set("")
    var_1.set("")
    var_2.set("")
    var_3.set("")
    var_4.set("")
    var_5.set("")
    var_6.set("")
    var_7.set("")
    var_8.set("")
    var_9.set("")
    var_10.set("")
    var_11.set("")
    var_12.set("")
    var_13.set("")
    summary_var.set("")


def test_variables(set, *args):
    if set == "simple essay":
        essay_name_var.set("The Amazon Rainforest")
        thesis_var.set("The Amazon Rainforest is important to humankind.")
        topic_sentence_1_var.set("The Amazon Rainforest is important for climate stability.")
        topic_sentence_2_var.set("The Amazon Rainforest provides a home for a diverse species.")
        topic_sentence_3_var.set("The Amazon Rainforest is beautiful and important for human wellbeing.")
        var_1.set("The Amazon Rainforest stores 2 billion metric tons of carbon annually.")
        var_2.set("The Amazon stores 200 billion metric tons of carbon in total.")
        var_3.set("The Amazon Rainforest has over ten of thousands of plant, animal and fungi species.")
        var_4.set("The Amazon has many unknown species and potential for new foods, medicines and products.")
        var_5.set("The Amazon Rainforest is home to over 30 million people.")
        var_6.set("The Amazon Rainforest is visited by one million people each year due to its beauty.")
        summary_var.set("The Amazon Rainforest is important for humankind and the entire planet.")

    elif set == "simple biography":
        essay_name_var.set("Ben Franklin")
        thesis_var.set("Ben Franklin was a American hero.")
        topic_sentence_1_var.set("He was a Founding Father of the United States.")
        topic_sentence_2_var.set("He was the first postmaster general of the United States.")
        topic_sentence_3_var.set("He secured strong support for the US Revolution from France as an important diplomat.")
        var_1.set("Boston, Massachusetts Bay, British America")
        var_2.set("January 6, 1705")
        var_3.set("Philadelphia, Pennsylvania, U.S.")
        var_4.set("April 17, 1790 (84)")
        var_5.set("Politician, Diplomat, Inventor, Scientist, Writer, Printer")
        var_6.set("6th President of Pennsylvania. Founding Father. First Postmaster General.")

    elif set == "simple resume":
        essay_name_var.set("Lloyd Andersen")
        thesis_var.set("702-589-8952")
        topic_sentence_1_var.set("Teamwork")
        topic_sentence_2_var.set("Leadership")
        topic_sentence_3_var.set("Communication")
        var_1.set("University of Nevada, Las Vegas")
        var_2.set("Bachelors of Science in Business, Finance, Cum Laude")
        var_3.set("Northwest Career and Technical Academy")
        var_4.set("Culinary Arts, Advanced Honors Diploma")
        var_5.set("Pizza Hut, 2020-2021")
        var_6.set("Head Cook, Shift Manager, Local Trainer")
        var_7.set("Pumpkin Patch, 2017, 2018")
        var_8.set("Arcade Station, Ride Operator")
        var_9.set("lloydandersenco@gmail.com")
        var_10.set("Python Language: Tkapp, WriteAlign")
        var_11.set("Microsoft Office: Word, Excel, Powerpoint")
        var_12.set("Content Writing: Code documentation, HomeSkool Lessons")
        var_13.set("Las Vegas, Nevada")
        summary_var.set("I am a enthusiastic team member. Whether in the kitchen or on a group portfolio optimization assignment. I am a practitioner of servant leadership and am great with virtual teams.")


    elif set == "simple cover letter":
        essay_name_var.set("Joe's Stock Yard")
        thesis_var.set("Stocking Manager")
        topic_sentence_1_var.set("Efficiency")
        topic_sentence_2_var.set("Organization")
        topic_sentence_3_var.set("Teamwork")
        var_1.set("Lewiston, Maine")
        var_2.set("Lloyd Andersen")
        var_3.set("Trainer")
        var_4.set("Pizza Hut")
        var_5.set("University of Las Vegas, Nevada")
        var_6.set("Bachelor of Science in Business, Finance, Cum Laude")
        summary_var.set("Las Vegas, Nevada")


def generate_simple_essay_prompt():
    file = open(f"{file_name_var.get()}.txt", "w")
    file.write(f"""Write me a five-paragraph essay.
The essay shall be constructed with an introduction paragraph, three body paragraphs and a concluding paragraph.

The title of the paper is {essay_name_var.get()}.

The thesis of the paper is {thesis_var.get()}.

The introduction paragraph shall utilize the body paragraphs for a coherent essay.
 
The body paragraphs should be written as follows:

The first body paragraph's topic sentence is "{topic_sentence_1_var.get()}"
It will have two supporting ideas as follows: 
1. {var_1.get()}
2. {var_2.get()}

The second body paragraph's topic sentence is "{topic_sentence_2_var.get()}"
It will have two supporting ideas as follows: 
1. {var_3.get()}
2. {var_4.get()}

The third body paragraph's topic sentence is "{topic_sentence_3_var.get()}"
It will have two supporting ideas as follows: 
1. {var_5.get()}
2. {var_6.get()}

The concluding paragraph shall summarize the body paragraphs.
The concluding paragraph shall emphasize the following statement: {summary_var.get()}
    """)
    file.close()

def generate_simple_resume_prompt():
    file = open(f"{file_name_var.get()}.txt", "w")
    file.write(f"""Write me a simple resume for my new job.
Make sure to organize the resume into groups.

My name is {essay_name_var.get()}.
My address is {var_13.get()}.
My phone number is {thesis_var.get()}.
My email address is {var_9.get()}.

There should be an about me section.
It should include the statement: {summary_var.get()}

My education background includes:
1. A {var_2.get()} from the {var_1.get()}.
2. A {var_4.get()} from the {var_3.get()}.

My work experience includes:
1. {var_6.get()} at {var_5.get()}.
2. {var_8.get()} at {var_7.get()}.

My top three traits are:
1. {topic_sentence_1_var.get()}
2. {topic_sentence_2_var.get()}
3. {topic_sentence_3_var.get()}

My top three skills are:
1. {var_10.get()}
2. {var_11.get()}
3. {var_12.get()}

Make sure to include all the information above in the resume.
    """)
    file.close()


def generate_simple_cover_letter_prompt():
    file = open(f"{file_name_var.get()}.txt", "w")
    file.write(f"""Please write a cover letter for a job I am trying to get.
The companies name is {essay_name_var.get()}.
{essay_name_var.get()} is located in {var_1.get()}.
The job title is called {thesis_var.get()}.
My name is {var_2.get()}.
I previously worked as a {var_3.get()} at {var_4.get()}.
I went to school at {var_5.get()} and got a {var_6.get()}.

The words I want to emphasize about the job are {topic_sentence_1_var.get()}, {topic_sentence_2_var.get()}, and {topic_sentence_3_var.get()}.

Make a case for why my previous work and education experience will help me be a good choice for the {thesis_var.get()} role.
Please limit the cover letter to three paragraphs.
    """)
    file.close()


def generate_simple_biography_prompt():
    file = open(f"{file_name_var.get()}.txt", "w")
    file.write(f"""Write be a biography about {essay_name_var.get()}
It should be a three paragraph narrative about their life.
It should begin with their birth and end with their death.
The thesis of the paper should be: {thesis_var.get()}
The three main points should be as follows:
1. {topic_sentence_1_var.get()}
2. {topic_sentence_2_var.get()}
3. {topic_sentence_3_var.get()}
The essay shall also include the following facts:
{essay_name_var.get()} were born in {var_1.get()} on {var_2.get()}.
{essay_name_var.get()} died in {var_3.get()} on {var_4.get()}.
{essay_name_var.get()} worked as {var_5.get()}.
Some major achievements of {essay_name_var.get()} include {var_6.get()}.
    """)
    file.close()


def generate_simple_essay_outline():
    document = Document()
    document.add_heading(f"{essay_name_var.get()}", 0)

    document.add_heading("Introduction", level=1)
    document.add_heading(f"-{thesis_var.get()}", level=2)

    document.add_heading("Body Paragraph 1", level=1)
    document.add_heading(f"-{topic_sentence_1_var.get()}", level=2)
    document.add_paragraph(f"-{var_1.get()}")
    document.add_paragraph(f"-{var_2.get()}")

    document.add_heading("Body Paragraph 2", level=1)
    document.add_heading(f"-{topic_sentence_2_var.get()}", level=2)
    document.add_paragraph(f"-{var_3.get()}")
    document.add_paragraph(f"-{var_4.get()}")

    document.add_heading("Body Paragraph 3", level=1)
    document.add_heading(f"-{topic_sentence_3_var.get()}", level=2)
    document.add_paragraph(f"-{var_5.get()}")
    document.add_paragraph(f"-{var_6.get()}")

    document.add_heading("Conclusion", level=1)
    document.add_heading(f"-{summary_var.get()}", level=2)

    document.save(f"{file_name_var.get()}.docx")


def generate_simple_resume_outline():
    document = Document()
    document.add_heading(f"{essay_name_var.get()}", level=0)
    document.add_heading("About me", level=1)
    document.add_paragraph(f"{essay_name_var.get()}")
    document.add_paragraph(f"{var_13.get()}")
    document.add_paragraph(f"{thesis_var.get()}")
    document.add_paragraph(f"{var_9.get()}")
    document.add_paragraph(f"{topic_sentence_1_var.get()}, {topic_sentence_2_var.get()}, {topic_sentence_3_var.get()}")
    document.add_paragraph(f"{summary_var.get()}")

    document.add_heading("Work Experience", level=1)
    document.add_paragraph(f"{var_5.get()}")
    document.add_paragraph(f"{var_6.get()}")
    document.add_paragraph(f"{var_7.get()}")
    document.add_paragraph(f"{var_8.get()}")

    document.add_heading("Educational Background", level=1)
    document.add_paragraph(f"{var_1.get()}")
    document.add_paragraph(f"{var_2.get()}")
    document.add_paragraph(f"{var_3.get()}")
    document.add_paragraph(f"{var_4.get()}")

    document.add_heading("Skills", level=1)
    document.add_paragraph(f"{var_10.get()}")
    document.add_paragraph(f"{var_11.get()}")
    document.add_paragraph(f"{var_12.get()}")

    document.save(f"{file_name_var.get()}.docx")


def generate_simple_cover_letter_outline():
    document = Document()
    # Your Name, Your Address, Today's date
    document.add_paragraph(f"{var_2.get()}")
    document.add_paragraph(f"{summary_var.get()}")
    document.add_paragraph(f"{dt.date.today()}")

    # Company Name and Address
    document.add_paragraph(f"{essay_name_var.get()}")
    document.add_paragraph(f"{var_1.get()}")

    document.add_paragraph(f"Dear {essay_name_var.get()}: (add hiring manager's name here)")
    document.add_paragraph(f"It is an amazing opportunity that you have the {thesis_var.get()} role open.")
    document.add_paragraph(f"As an aspiring {thesis_var.get()}, my educational and work history was made for this opportunity.")
    document.add_paragraph(f"I heard about {essay_name_var.get()} from the career fair at my school.")
    document.add_paragraph(f"And ever since I have been excited about it.")

    document.add_paragraph(f"I hope to bring {topic_sentence_1_var.get()}, {topic_sentence_2_var.get()} and {topic_sentence_3_var.get()} to this role.")

    document.add_paragraph(f"My educational background includes a {var_6.get()} from the {var_5.get()}.")
    document.add_paragraph(f"In my past roles I learned alot about  {topic_sentence_1_var.get()}, {topic_sentence_2_var.get()} and {topic_sentence_3_var.get()}.")
    document.add_paragraph(f"I worked for {var_4.get()} as a {var_3.get()}. The above skills were very important for my past roles.")
    document.add_paragraph(f"Thank You for this opportunity,")
    document.add_paragraph(f"{var_2.get()}")
    document.add_paragraph(f"(Digital Signature)")

    document.add_paragraph(f"Resume Enclosed.")
    document.save(f"{file_name_var.get()}.docx")




def generate_simple_biography_outline():
    document = Document()
    document.add_heading(f"{essay_name_var.get()}", 0)
    document.add_heading("Life Details", level=1)
    document.add_heading(f"Born in {var_1.get()} on {var_2.get()}", level=2)
    document.add_heading(f"Died in {var_3.get()} on {var_4.get()}", level=2)
    document.add_heading(f"Work: {var_5.get()}", level=2)
    document.add_heading(f"Achievements: {var_6.get()}", level=2)
    document.add_heading(f"{thesis_var.get()}", level=1)
    document.add_heading(f"{topic_sentence_1_var.get()}", level=2)
    document.add_heading(f"{topic_sentence_2_var.get()}", level=2)
    document.add_heading(f"{topic_sentence_3_var.get()}", level=2)

    document.save(f"{file_name_var.get()}.docx")


def simple_essay_prompt():
    top = tk.Toplevel(background="white")

    file_name_label = tk.Label(top, text="File Name", background="white")
    file_name_label.grid(row=0, column=0)

    file_name_entry = tk.Entry(top, textvariable=file_name_var)
    file_name_entry.grid(row=0, column=1)

    file_type_label = tk.Label(top, text=".txt", background="white")
    file_type_label.grid(row=0, column=2)

    create_button = tk.Button(top, text="Create", background="white", command=generate_simple_essay_prompt)
    create_button.grid(row=1, column=0, columnspan=3, sticky="swen")


def simple_essay_outline():
    top = tk.Toplevel(background="white")

    file_name_label = tk.Label(top, text="File Name", background="white")
    file_name_label.grid(row=0, column=0)

    file_name_entry = tk.Entry(top, textvariable=file_name_var)
    file_name_entry.grid(row=0, column=1)

    file_type_label = tk.Label(top, text=".docx", background="white")
    file_type_label.grid(row=0, column=2)

    create_button = tk.Button(top, text="Create", background="white", command=generate_simple_essay_outline)
    create_button.grid(row=1, column=0, columnspan=3, sticky="swen")


def simple_resume_outline():
    top = tk.Toplevel(background="white")

    file_name_label = tk.Label(top, text="File Name", background="white")
    file_name_label.grid(row=0, column=0)

    file_name_entry = tk.Entry(top, textvariable=file_name_var)
    file_name_entry.grid(row=0, column=1)

    file_type_label = tk.Label(top, text=".docx", background="white")
    file_type_label.grid(row=0, column=2)

    create_button = tk.Button(top, text="Create", background="white", command=generate_simple_resume_outline)
    create_button.grid(row=1, column=0, columnspan=3, sticky="swen")


def simple_resume_prompt():
    top = tk.Toplevel(background="white")

    file_name_label = tk.Label(top, text="File Name", background="white")
    file_name_label.grid(row=0, column=0)

    file_name_entry = tk.Entry(top, textvariable=file_name_var)
    file_name_entry.grid(row=0, column=1)

    file_type_label = tk.Label(top, text=".txt", background="white")
    file_type_label.grid(row=0, column=2)

    create_button = tk.Button(top, text="Create", background="white", command=generate_simple_resume_prompt)
    create_button.grid(row=1, column=0, columnspan=3, sticky="swen")


def simple_biography_outline():
    top = tk.Toplevel(background="white")

    file_name_label = tk.Label(top, text="File Name", background="white")
    file_name_label.grid(row=0, column=0)

    file_name_entry = tk.Entry(top, textvariable=file_name_var)
    file_name_entry.grid(row=0, column=1)

    file_type_label = tk.Label(top, text=".docx", background="white")
    file_type_label.grid(row=0, column=2)

    create_button = tk.Button(top, text="Create", background="white", command=generate_simple_biography_outline)
    create_button.grid(row=1, column=0, columnspan=3, sticky="swen")



def simple_biography_prompt():
    top = tk.Toplevel(background="white")

    file_name_label = tk.Label(top, text="File Name", background="white")
    file_name_label.grid(row=0, column=0)

    file_name_entry = tk.Entry(top, textvariable=file_name_var)
    file_name_entry.grid(row=0, column=1)

    file_type_label = tk.Label(top, text=".txt", background="white")
    file_type_label.grid(row=0, column=2)

    create_button = tk.Button(top, text="Create", background="white", command=generate_simple_biography_prompt)
    create_button.grid(row=1, column=0, columnspan=3, sticky="swen")


def simple_cover_letter_outline():
    top = tk.Toplevel(background="white")

    file_name_label = tk.Label(top, text="File Name", background="white")
    file_name_label.grid(row=0, column=0)

    file_name_entry = tk.Entry(top, textvariable=file_name_var)
    file_name_entry.grid(row=0, column=1)

    file_type_label = tk.Label(top, text=".docx", background="white")
    file_type_label.grid(row=0, column=2)

    create_button = tk.Button(top, text="Create", background="white", command=generate_simple_cover_letter_outline)
    create_button.grid(row=1, column=0, columnspan=3, sticky="swen")


def simple_cover_letter_prompt():
    top = tk.Toplevel(background="white")

    file_name_label = tk.Label(top, text="File Name", background="white")
    file_name_label.grid(row=0, column=0)

    file_name_entry = tk.Entry(top, textvariable=file_name_var)
    file_name_entry.grid(row=0, column=1)

    file_type_label = tk.Label(top, text=".txt", background="white")
    file_type_label.grid(row=0, column=2)

    create_button = tk.Button(top, text="Create", background="white", command=generate_simple_cover_letter_prompt)
    create_button.grid(row=1, column=0, columnspan=3, sticky="swen")


def back_to_home():
    for child in root.winfo_children():
        child.pack_forget()
    main_page.pack(fill="both", side="top", padx=150, pady=150)


def clear_general_frame():
    for child in general_frame.winfo_children():
        child.destroy()


def simple_essay():
    simple_essay_frame = tk.Frame(root, background="white")
    simple_essay_frame.pack(side="top", pady=(10, 10), padx=10)
    simple_essay_frame.columnconfigure((0, 1), weight=1)

    back_button = tk.Button(simple_essay_frame, text="back", background="black", foreground="white", command=back_to_home)
    back_button.grid(row=0, column=0)

    simple_essay_title = tk.Label(simple_essay_frame, text="Simple Essay", background="white", foreground="black", font=("Times New Roman", 30, "bold"))
    simple_essay_title.grid(row=1, column=1, columnspan=2, pady=(0, 30))


    simple_essay_name_label = tk.Label(simple_essay_frame, text="Essay Name", background="white", foreground="black", font=label_font)
    simple_essay_name_label.grid(row=2, column=1, pady=(0, 5), padx=(0, 5))

    simple_essay_name_entry = tk.Entry(simple_essay_frame, textvariable=essay_name_var,  relief="solid", font=entry_font)
    simple_essay_name_entry.grid(row=2, column=2, ipadx=80)

    simple_essay_thesis_label = tk.Label(simple_essay_frame, text="Thesis", background="white", foreground="black", font=label_font)
    simple_essay_thesis_label.grid(row=3, column=1, pady=(0, 5))

    simple_essay_thesis_entry = tk.Entry(simple_essay_frame, textvariable=thesis_var, relief="solid", font=entry_font)
    simple_essay_thesis_entry.grid(row=3, column=2, ipadx=80)

    topic_1_label = tk.Label(simple_essay_frame, text="Topic 1", background="white", foreground="black", font=label_font)
    topic_1_label.grid(row=4, column=1, pady=(0, 5))

    topic_1_entry = tk.Entry(simple_essay_frame, textvariable=topic_sentence_1_var, relief="solid", font=entry_font)
    topic_1_entry.grid(row=4, column=2, ipadx=80)

    support_1_label = tk.Label(simple_essay_frame, text="Support 1", background="white", foreground="black", font=label_font)
    support_1_label.grid(row=5, column=1, pady=(0, 5))

    support_1_entry = tk.Entry(simple_essay_frame, textvariable=var_1, relief="solid", font=entry_font)
    support_1_entry.grid(row=5, column=2, ipadx=80)

    support_2_label = tk.Label(simple_essay_frame, text="Support 2", background="white", foreground="black", font=label_font)
    support_2_label.grid(row=6, column=1, pady=(0, 5))

    support_2_entry = tk.Entry(simple_essay_frame, textvariable=var_2, relief="solid", font=entry_font)
    support_2_entry.grid(row=6, column=2, ipadx=80)

    topic_2_label = tk.Label(simple_essay_frame, text="Topic 2", background="white", foreground="black", font=label_font)
    topic_2_label.grid(row=7, column=1, pady=(0, 5))

    topic_2_entry = tk.Entry(simple_essay_frame, textvariable=topic_sentence_2_var, relief="solid", font=entry_font)
    topic_2_entry.grid(row=7, column=2, ipadx=80)

    support_3_label = tk.Label(simple_essay_frame, text="Support 1", background="white", foreground="black", font=label_font)
    support_3_label.grid(row=8, column=1, pady=(0, 5))

    support_3_entry = tk.Entry(simple_essay_frame, textvariable=var_3, relief="solid", font=entry_font)
    support_3_entry.grid(row=8, column=2, ipadx=80)

    support_4_label = tk.Label(simple_essay_frame, text="Support 2", background="white", foreground="black", font=label_font)
    support_4_label.grid(row=9, column=1, pady=(0, 5))

    support_4_entry = tk.Entry(simple_essay_frame, textvariable=var_4, relief="solid", font=entry_font)
    support_4_entry.grid(row=9, column=2, ipadx=80)

    topic_3_label = tk.Label(simple_essay_frame, text="Topic 3", background="white", foreground="black", font=label_font)
    topic_3_label.grid(row=10, column=1, pady=(0, 5))

    topic_3_entry = tk.Entry(simple_essay_frame, textvariable=topic_sentence_3_var, relief="solid", font=entry_font)
    topic_3_entry.grid(row=10, column=2, ipadx=80)

    support_5_label = tk.Label(simple_essay_frame, text="Support 1", background="white", foreground="black", font=label_font)
    support_5_label.grid(row=11, column=1, pady=(0, 5))

    support_5_entry = tk.Entry(simple_essay_frame, textvariable=var_5, relief="solid", font=entry_font)
    support_5_entry.grid(row=11, column=2, ipadx=80)

    support_6_label = tk.Label(simple_essay_frame, text="Support 2", background="white", foreground="black", font=label_font)
    support_6_label.grid(row=12, column=1, pady=(0, 5))

    support_6_entry = tk.Entry(simple_essay_frame, textvariable=var_6, relief="solid", font=entry_font)
    support_6_entry.grid(row=12, column=2, ipadx=80)


    finish_frame = tk.Frame(simple_essay_frame, background="white")
    finish_frame.grid(row=13, column=1, columnspan=2)

    prompt_generation_button = tk.Button(finish_frame, text="Prompt", background="white", foreground="black", command=simple_essay_prompt)
    prompt_generation_button.grid(row=0, column=0, ipadx=20, ipady=10)

    outline_generation_button = tk.Button(finish_frame, text="Outline", background="white", foreground="black", command=simple_essay_outline)
    outline_generation_button.grid(row=0, column=1, ipadx=20, ipady=10)

    clear_button = tk.Button(finish_frame, text="Clear", background="black", foreground="white")
    clear_button.grid(row=1, column=0, columnspan=2, sticky="swen")
    clear_button.bind("<Double-Button-1>", reset_variables)
    clear_button.bind("<Button-3>", lambda e: test_variables('simple essay'))

def simple_resume_2():

    clear_general_frame()

    for child in root.winfo_children():
        child.pack_forget()


    general_frame = tk.Frame(root, background="white")
    general_frame.pack(side="top", pady=(10, 10), padx=10)
    general_frame.columnconfigure((0, 1), weight=1)

    back_button = tk.Button(general_frame, text="back", background="black", foreground="white",
                            command=back_to_home)
    back_button.grid(row=0, column=0)

    simple_essay_title = tk.Label(general_frame, text="Simple Resume", background="white", foreground="black",
                                  font=("Times New Roman", 30, "bold"))
    simple_essay_title.grid(row=1, column=1, columnspan=2, pady=(0, 30))

    skill_one_label = tk.Label(general_frame, text="Skill 1", background="white", foreground="black",
                               font=label_font)
    skill_one_label.grid(row=2, column=1, pady=(0, 5), padx=(0, 5))

    skill_one_entry = tk.Entry(general_frame, textvariable=var_10, relief="solid", font=entry_font)
    skill_one_entry.grid(row=2, column=2, ipadx=80)

    skill_two_label = tk.Label(general_frame, text="Skill 2", background="white", foreground="black",
                               font=label_font)
    skill_two_label.grid(row=3, column=1, pady=(0, 5), padx=(0, 5))

    skill_two_entry = tk.Entry(general_frame, textvariable=var_11, relief="solid", font=entry_font)
    skill_two_entry.grid(row=3, column=2, ipadx=80)

    skill_three_label = tk.Label(general_frame, text="Skill 3", background="white", foreground="black",
                                 font=label_font)
    skill_three_label.grid(row=4, column=1, pady=(0, 5), padx=(0, 5))

    skill_three_entry = tk.Entry(general_frame, textvariable=var_12, relief="solid", font=entry_font)
    skill_three_entry.grid(row=4, column=2, ipadx=80)

    point_1_label = tk.Label(general_frame, text="Point 1", background="white", foreground="black",
                             font=label_font)
    point_1_label.grid(row=5, column=1, pady=(0, 5), padx=(0, 5))

    point_1_entry = tk.Entry(general_frame, textvariable=topic_sentence_1_var, relief="solid", font=entry_font)
    point_1_entry.grid(row=5, column=2, ipadx=80)

    point_2_label = tk.Label(general_frame, text="Point 2", background="white", foreground="black",
                             font=label_font)
    point_2_label.grid(row=6, column=1, pady=(0, 5), padx=(0, 5))

    point_2_entry = tk.Entry(general_frame, textvariable=topic_sentence_2_var, relief="solid", font=entry_font)
    point_2_entry.grid(row=6, column=2, ipadx=80)

    point_3_label = tk.Label(general_frame, text="Point 3", background="white", foreground="black",
                             font=label_font)
    point_3_label.grid(row=7, column=1, padx=(0, 5))

    point_3_entry = tk.Entry(general_frame, textvariable=topic_sentence_3_var, relief="solid", font=entry_font)
    point_3_entry.grid(row=7, column=2, ipadx=80)

    about_me_label = tk.Label(general_frame, text="About Me", background="white", foreground="black",
                             font=label_font)
    about_me_label.grid(row=8, column=1, padx=(0, 5))

    about_me_entry = tk.Entry(general_frame, textvariable=summary_var, relief="solid", font=entry_font)
    about_me_entry.grid(row=8, column=2, ipadx=80)


    finish_frame = tk.Frame(general_frame, background="white")
    finish_frame.grid(row=9, column=1, columnspan=2)

    prompt_generation_button = tk.Button(finish_frame, text="Prompt", background="white", foreground="black",
                                         command=simple_resume_prompt)
    prompt_generation_button.grid(row=0, column=0, ipadx=20, ipady=10)

    outline_generation_button = tk.Button(finish_frame, text="Outline", background="white", foreground="black",
                                          command=simple_resume_outline)
    outline_generation_button.grid(row=0, column=1, ipadx=20, ipady=10)

    clear_button = tk.Button(finish_frame, text="Clear", background="black", foreground="white")
    clear_button.grid(row=1, column=0, columnspan=2, sticky="swen")
    clear_button.bind("<Double-Button-1>", reset_variables)
    clear_button.bind("<Button-3>", lambda e: test_variables('simple resume'))



def simple_resume():
    clear_general_frame()

    general_frame = tk.Frame(root, background="white")
    general_frame.pack(side="top", pady=(10, 10), padx=10)
    general_frame.columnconfigure((0, 1), weight=1)

    back_button = tk.Button(general_frame, text="back", background="black", foreground="white",
                            command=back_to_home)
    back_button.grid(row=0, column=0)

    simple_essay_title = tk.Label(general_frame, text="Simple Resume", background="white", foreground="black",
                                  font=("Times New Roman", 30, "bold"))
    simple_essay_title.grid(row=1, column=1, columnspan=2, pady=(0, 30))

    # sep
    simple_resume_name_label = tk.Label(general_frame, text="Name", background="white", foreground="black",
                                       font=label_font)
    simple_resume_name_label.grid(row=2, column=1, pady=(0, 5), padx=(0, 5))

    simple_resume_name_entry = tk.Entry(general_frame, textvariable=essay_name_var, relief="solid", font=entry_font)
    simple_resume_name_entry.grid(row=2, column=2, ipadx=80)

    highest_education_name_label = tk.Label(general_frame, text="School Name", background="white", foreground="black",
                                    font=label_font)
    highest_education_name_label.grid(row=3, column=1, pady=(0, 5), padx=(0, 5))

    highest_education_name_entry = tk.Entry(general_frame, textvariable=var_1, relief="solid", font=entry_font)
    highest_education_name_entry.grid(row=3, column=2, ipadx=80)

    highest_degree_label = tk.Label(general_frame, text="Degree", background="white", foreground="black",
                                font=label_font)
    highest_degree_label.grid(row=4, column=1, pady=(0, 5), padx=(0, 5))

    highest_degree_entry = tk.Entry(general_frame, textvariable=var_2, relief="solid", font=entry_font)
    highest_degree_entry.grid(row=4, column=2, ipadx=80)

    # Education

    secondary_education_label = tk.Label(general_frame, text="School Name", background="white", foreground="black",
                                    font=label_font)
    secondary_education_label.grid(row=5, column=1, pady=(0, 5), padx=(0, 5))

    secondary_education_entry = tk.Entry(general_frame, textvariable=var_3, relief="solid", font=entry_font)
    secondary_education_entry.grid(row=5, column=2, ipadx=80)

    secondary_education_degree_label = tk.Label(general_frame, text="Degree", background="white", foreground="black",
                                font=label_font)
    secondary_education_degree_label.grid(row=6, column=1, pady=(0, 5), padx=(0, 5))

    secondary_education_degree_entry = tk.Entry(general_frame, textvariable=var_4, relief="solid", font=entry_font)
    secondary_education_degree_entry.grid(row=6, column=2, ipadx=80)

    # Work, Achievements
    last_job_label = tk.Label(general_frame, text="Job", background="white", foreground="black",
                         font=label_font)
    last_job_label.grid(row=7, column=1, pady=(0, 5), padx=(0, 5))

    last_job_entry = tk.Entry(general_frame, textvariable=var_5, relief="solid", font=entry_font)
    last_job_entry.grid(row=7, column=2, ipadx=80)

    last_job_role_label = tk.Label(general_frame, text="Role", background="white", foreground="black",
                                  font=label_font)
    last_job_role_label.grid(row=8, column=1, pady=(0, 5), padx=(0, 5))

    last_job_role_entry = tk.Entry(general_frame, textvariable=var_6, relief="solid", font=entry_font)
    last_job_role_entry.grid(row=8, column=2, ipadx=80)
    #

    old_job_label = tk.Label(general_frame, text="Job", background="white", foreground="black",
                              font=label_font)
    old_job_label.grid(row=9, column=1, pady=(0, 5), padx=(0, 5))

    old_job_entry = tk.Entry(general_frame, textvariable=var_7, relief="solid", font=entry_font)
    old_job_entry.grid(row=9, column=2, ipadx=80)

    old_job_role_label = tk.Label(general_frame, text="Role", background="white", foreground="black",
                                   font=label_font)
    old_job_role_label.grid(row=10, column=1, pady=(0, 5), padx=(0, 5))

    old_job_role_entry = tk.Entry(general_frame, textvariable=var_8, relief="solid", font=entry_font)
    old_job_role_entry.grid(row=10, column=2, ipadx=80)

    # Contact

    phone_number_label = tk.Label(general_frame, text="Number", background="white", foreground="black",
                            font=label_font)
    phone_number_label.grid(row=11, column=1, pady=(0, 5), padx=(0, 5))

    phone_number_entry = tk.Entry(general_frame, textvariable=thesis_var, relief="solid", font=entry_font)
    phone_number_entry.grid(row=11, column=2, ipadx=80)

    email_label = tk.Label(general_frame, text="Email", background="white", foreground="black",
                                  font=label_font)
    email_label.grid(row=12, column=1, pady=(0, 5), padx=(0, 5))

    email_entry = tk.Entry(general_frame, textvariable=var_9, relief="solid", font=entry_font)
    email_entry.grid(row=12, column=2, ipadx=80)

    address_label = tk.Label(general_frame, text="Address", background="white", foreground="black",
                           font=label_font)
    address_label.grid(row=13, column=1, pady=(0, 5), padx=(0, 5))

    address_entry = tk.Entry(general_frame, textvariable=var_13, relief="solid", font=entry_font)
    address_entry.grid(row=13, column=2, ipadx=80)

    next_page_button = tk.Button(general_frame, text="Next", command=simple_resume_2, background="white", foreground="black")
    next_page_button.grid(row=14, column=1, columnspan=2, pady=(20, 0), ipadx=80)



    # Finish
    finish_frame = tk.Frame(general_frame, background="white")
    finish_frame.grid(row=15, column=1, columnspan=2)

    clear_button = tk.Button(finish_frame, text="Clear", background="black", foreground="white")
    clear_button.grid(row=0, column=0, columnspan=2, sticky="swen", ipadx=80)
    clear_button.bind("<Double-Button-1>", reset_variables)
    clear_button.bind("<Button-3>", lambda e: test_variables('simple resume'))


def simple_cover_letter():
    clear_general_frame()
    general_frame = tk.Frame(root, background="white")
    general_frame.pack(side="top", pady=(10, 10), padx=10)
    general_frame.columnconfigure((0, 1), weight=1)

    back_button = tk.Button(general_frame, text="back", background="black", foreground="white",
                            command=back_to_home)
    back_button.grid(row=0, column=0)

    simple_essay_title = tk.Label(general_frame, text="Simple Cover Letter", background="white", foreground="black",
                                  font=("Times New Roman", 30, "bold"))
    simple_essay_title.grid(row=1, column=1, columnspan=2, pady=(0, 30))

    # Entries
    company_name_label = tk.Label(general_frame, text="Company", background="white", foreground="black",
                                        font=label_font)
    company_name_label.grid(row=2, column=1, pady=(0, 5), padx=(0, 5))

    company_name_entry = tk.Entry(general_frame, textvariable=essay_name_var, relief="solid", font=entry_font)
    company_name_entry.grid(row=2, column=2, ipadx=80)

    company_location_label = tk.Label(general_frame, text="Location", background="white", foreground="black",
                                  font=label_font)
    company_location_label.grid(row=3, column=1, pady=(0, 5), padx=(0, 5))

    company_location_entry = tk.Entry(general_frame, textvariable=var_1, relief="solid", font=entry_font)
    company_location_entry.grid(row=3, column=2, ipadx=80)

    role_name_label = tk.Label(general_frame, text="Role Title", background="white", foreground="black",
                                  font=label_font)
    role_name_label.grid(row=4, column=1, pady=(0, 5), padx=(0, 5))

    role_name_entry = tk.Entry(general_frame, textvariable=thesis_var, relief="solid", font=entry_font)
    role_name_entry.grid(row=4, column=2, ipadx=80)

    name_label = tk.Label(general_frame, text="Name", background="white", foreground="black",
                                  font=label_font)
    name_label.grid(row=5, column=1, pady=(0, 5), padx=(0, 5))

    name_entry = tk.Entry(general_frame, textvariable=var_2, relief="solid", font=entry_font)
    name_entry.grid(row=5, column=2, ipadx=80)

    address_label = tk.Label(general_frame, text="Address", background="white", foreground="black",
                                  font=label_font)
    address_label.grid(row=6, column=1, pady=(0, 5), padx=(0, 5))

    address_entry = tk.Entry(general_frame, textvariable=summary_var, relief="solid", font=entry_font)
    address_entry.grid(row=6, column=2, ipadx=80)

    last_company_name_label = tk.Label(general_frame, text="Last Job Name", background="white", foreground="black",
                                  font=label_font)
    last_company_name_label.grid(row=7, column=1, pady=(0, 5), padx=(0, 5))

    last_company_name_entry = tk.Entry(general_frame, textvariable=var_4, relief="solid", font=entry_font)
    last_company_name_entry.grid(row=7, column=2, ipadx=80)

    last_company_title_label = tk.Label(general_frame, text="Last Job Title", background="white", foreground="black",
                                  font=label_font)
    last_company_title_label.grid(row=8, column=1, pady=(0, 5), padx=(0, 5))

    last_company_title_entry = tk.Entry(general_frame, textvariable=var_3, relief="solid", font=entry_font)
    last_company_title_entry.grid(row=8, column=2, ipadx=80)

    school_name_label = tk.Label(general_frame, text="School", background="white", foreground="black",
                                  font=label_font)
    school_name_label.grid(row=9, column=1, pady=(0, 5), padx=(0, 5))

    school_name_entry = tk.Entry(general_frame, textvariable=var_5, relief="solid", font=entry_font)
    school_name_entry.grid(row=9, column=2, ipadx=80)
    degree_name_label = tk.Label(general_frame, text="Degree", background="white", foreground="black",
                                  font=label_font)
    degree_name_label.grid(row=10, column=1, pady=(0, 5), padx=(0, 5))

    degree_name_entry = tk.Entry(general_frame, textvariable=var_6, relief="solid", font=entry_font)
    degree_name_entry.grid(row=10, column=2, ipadx=80)

    key_one_label = tk.Label(general_frame, text="Key 1", background="white", foreground="black",
                                  font=label_font)
    key_one_label.grid(row=11, column=1, pady=(0, 5), padx=(0, 5))

    key_one_entry = tk.Entry(general_frame, textvariable=topic_sentence_1_var, relief="solid", font=entry_font)
    key_one_entry.grid(row=11, column=2, ipadx=80)

    key_two_label = tk.Label(general_frame, text="Key 2", background="white", foreground="black",
                                  font=label_font)
    key_two_label.grid(row=12, column=1, pady=(0, 5), padx=(0, 5))

    key_two_entry = tk.Entry(general_frame, textvariable=topic_sentence_2_var, relief="solid", font=entry_font)
    key_two_entry.grid(row=12, column=2, ipadx=80)

    key_three_label = tk.Label(general_frame, text="Key 3", background="white", foreground="black",
                                  font=label_font)
    key_three_label.grid(row=13, column=1, pady=(0, 5), padx=(0, 5))

    key_three_entry = tk.Entry(general_frame, textvariable=topic_sentence_3_var, relief="solid", font=entry_font)
    key_three_entry.grid(row=13, column=2, ipadx=80)



    finish_frame = tk.Frame(general_frame, background="white")
    finish_frame.grid(row=14, column=1, columnspan=2)

    prompt_generation_button = tk.Button(finish_frame, text="Prompt", background="white", foreground="black",
                                         command=simple_cover_letter_prompt)
    prompt_generation_button.grid(row=0, column=0, ipadx=20, ipady=10)

    outline_generation_button = tk.Button(finish_frame, text="Outline", background="white", foreground="black",
                                          command=simple_cover_letter_outline)
    outline_generation_button.grid(row=0, column=1, ipadx=20, ipady=10)

    clear_button = tk.Button(finish_frame, text="Clear", background="black", foreground="white")
    clear_button.grid(row=1, column=0, columnspan=2, sticky="swen")
    clear_button.bind("<Double-Button-1>", reset_variables)
    clear_button.bind("<Button-3>", lambda e: test_variables('simple cover letter'))



def simple_biography():
    clear_general_frame()
    general_frame = tk.Frame(root, background="white")
    general_frame.pack(side="top", pady=(10, 10), padx=10)
    general_frame.columnconfigure((0, 1), weight=1)

    back_button = tk.Button(general_frame, text="back", background="black", foreground="white",
                            command=back_to_home)
    back_button.grid(row=0, column=0)

    simple_essay_title = tk.Label(general_frame, text="Simple Biography", background="white", foreground="black",
                                  font=("Times New Roman", 30, "bold"))
    simple_essay_title.grid(row=1, column=1, columnspan=2, pady=(0, 30))

    simple_essay_name_label = tk.Label(general_frame, text="Name", background="white", foreground="black",
                                       font=label_font)
    simple_essay_name_label.grid(row=2, column=1, pady=(0, 5), padx=(0, 5))

    simple_essay_name_entry = tk.Entry(general_frame, textvariable=essay_name_var, relief="solid", font=entry_font)
    simple_essay_name_entry.grid(row=2, column=2, ipadx=80)

    birth_location_label = tk.Label(general_frame, text="Birth Place", background="white", foreground="black",
                                       font=label_font)
    birth_location_label.grid(row=3, column=1, pady=(0, 5), padx=(0, 5))

    birth_location_entry = tk.Entry(general_frame, textvariable=var_1, relief="solid", font=entry_font)
    birth_location_entry.grid(row=3, column=2, ipadx=80)

    birth_date_label = tk.Label(general_frame, text="Birth Date", background="white", foreground="black",
                                       font=label_font)
    birth_date_label.grid(row=4, column=1, pady=(0, 5), padx=(0, 5))

    birth_date_entry = tk.Entry(general_frame, textvariable=var_2, relief="solid", font=entry_font)
    birth_date_entry.grid(row=4, column=2, ipadx=80)

    # Death

    death_location_label = tk.Label(general_frame, text="Death Place", background="white", foreground="black",
                                    font=label_font)
    death_location_label.grid(row=5, column=1, pady=(0, 5), padx=(0, 5))

    death_location_entry = tk.Entry(general_frame, textvariable=var_3, relief="solid", font=entry_font)
    death_location_entry.grid(row=5, column=2, ipadx=80)

    death_date_label = tk.Label(general_frame, text="Death Date (age)", background="white", foreground="black",
                                font=label_font)
    death_date_label.grid(row=6, column=1, pady=(0, 5), padx=(0, 5))

    death_date_entry = tk.Entry(general_frame, textvariable=var_4, relief="solid", font=entry_font)
    death_date_entry.grid(row=6, column=2, ipadx=80)

    # Work, Achievements
    job_label = tk.Label(general_frame, text="Work", background="white", foreground="black",
                                    font=label_font)
    job_label.grid(row=7, column=1, pady=(0, 5), padx=(0, 5))

    job_entry = tk.Entry(general_frame, textvariable=var_5, relief="solid", font=entry_font)
    job_entry.grid(row=7, column=2, ipadx=80)

    achievements_label = tk.Label(general_frame, text="Achievements", background="white", foreground="black",
                                font=label_font)
    achievements_label.grid(row=8, column=1, pady=(0, 5), padx=(0, 5))

    achievements_entry = tk.Entry(general_frame, textvariable=var_6, relief="solid", font=entry_font)
    achievements_entry.grid(row=8, column=2, ipadx=80)

    # Narrative

    thesis_label = tk.Label(general_frame, text="Thesis", background="white", foreground="black",
                         font=label_font)
    thesis_label.grid(row=9, column=1, pady=(0, 5), padx=(0, 5))

    thesis_entry = tk.Entry(general_frame, textvariable=thesis_var, relief="solid", font=entry_font)
    thesis_entry.grid(row=9, column=2, ipadx=80)

    point_1_label = tk.Label(general_frame, text="Point 1", background="white", foreground="black",
                                  font=label_font)
    point_1_label.grid(row=10, column=1, pady=(0, 5), padx=(0, 5))

    point_1_entry = tk.Entry(general_frame, textvariable=topic_sentence_1_var, relief="solid", font=entry_font)
    point_1_entry.grid(row=10, column=2, ipadx=80)

    point_2_label = tk.Label(general_frame, text="Point 2", background="white", foreground="black",
                         font=label_font)
    point_2_label.grid(row=11, column=1, pady=(0, 5), padx=(0, 5))

    point_2_entry = tk.Entry(general_frame, textvariable=topic_sentence_2_var, relief="solid", font=entry_font)
    point_2_entry.grid(row=11, column=2, ipadx=80)

    point_3_label = tk.Label(general_frame, text="Point 3", background="white", foreground="black",
                                  font=label_font)
    point_3_label.grid(row=12, column=1, pady=(0, 5), padx=(0, 5))

    point_3_entry = tk.Entry(general_frame, textvariable=topic_sentence_3_var, relief="solid", font=entry_font)
    point_3_entry.grid(row=12, column=2, ipadx=80)


    finish_frame = tk.Frame(general_frame, background="white")
    finish_frame.grid(row=13, column=1, columnspan=2)

    prompt_generation_button = tk.Button(finish_frame, text="Prompt", background="white", foreground="black",
                                         command=simple_biography_prompt)
    prompt_generation_button.grid(row=0, column=0, ipadx=20, ipady=10)

    outline_generation_button = tk.Button(finish_frame, text="Outline", background="white", foreground="black",
                                          command=simple_biography_outline)
    outline_generation_button.grid(row=0, column=1, ipadx=20, ipady=10)

    clear_button = tk.Button(finish_frame, text="Clear", background="black", foreground="white")
    clear_button.grid(row=1, column=0, columnspan=2, sticky="swen")
    clear_button.bind("<Double-Button-1>", reset_variables)
    clear_button.bind("<Button-3>", lambda e: test_variables('simple biography'))




def make_selection():
    name = template_var.get()
    main_page.pack_forget()
    if name == "Simple Essay":
        simple_essay()
    elif name == "Simple Resume":
        simple_resume()
    elif name == "Simple Cover Letter":
        simple_cover_letter()
    elif name == "Simple Biography":
        simple_biography()
    else:
        main_page.pack(side="top", fill="both")


select_button = tk.Button(main_page, text="Write", background="white", foreground="black", font=("Times New Roman", 16), borderwidth=2, command=make_selection)
select_button.grid(row=2, column=0, pady=(10, 0), ipadx=50)

select_template.bind("<<ComboboxSelected>>",lambda e: main_page.focus())

root.mainloop()
    